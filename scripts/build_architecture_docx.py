from __future__ import annotations

import os
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "architecture" / "Canada_Logistics_Marketplace_Architecture_v1.0.docx"
QA_DIR = Path(os.environ.get("ARCH_DOC_QA_DIR", "/private/tmp/canada_marketplace_architecture"))
QA_DIR.mkdir(parents=True, exist_ok=True)

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "0B2545"
INK = "1F2937"
MUTED = "5F6B7A"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
PALE_BLUE = "F4F7FB"
WHITE = "FFFFFF"
BORDER = "CCD6E2"
GREEN = "2E6B50"
AMBER = "8A5A00"
RED = "9B1C1C"

BASE_FONT = "Calibri"
EAST_ASIA_FONT = "PingFang SC"
CODE_FONT = "Menlo"


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_run_font(run, *, name=BASE_FONT, size=None, color=INK, bold=None, italic=None):
    run.font.name = name
    if run._element.get_or_add_rPr().rFonts is None:
        run._element.get_or_add_rPr().append(OxmlElement("w:rFonts"))
    rfonts = run._element.get_or_add_rPr().rFonts
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:cs"), name)
    rfonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_border(cell, color=BORDER, size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "start", "bottom", "end", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        node = borders.find(tag)
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa: list[int], indent_dxa=120):
    if sum(widths_dxa) != 9360:
        raise ValueError(f"table widths must total 9360 DXA, got {sum(widths_dxa)}")
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    cell_mar = tbl_pr.find(qn("w:tblCellMar"))
    if cell_mar is None:
        cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(cell_mar)
    for edge, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
        node = cell_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            cell_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[i] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[i]))
            tc_w.set(qn("w:type"), "dxa")


def set_paragraph_keep(paragraph, *, keep_next=False, keep_lines=True):
    p_pr = paragraph._p.get_or_add_pPr()
    if keep_next:
        node = p_pr.find(qn("w:keepNext"))
        if node is None:
            node = OxmlElement("w:keepNext")
            p_pr.append(node)
    if keep_lines:
        node = p_pr.find(qn("w:keepLines"))
        if node is None:
            node = OxmlElement("w:keepLines")
            p_pr.append(node)


def add_numbering_definition(doc: Document, *, kind: str) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if kind == "bullet" else "%1.")
    lvl.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    lvl.append(lvl_jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.append(ind)
    lvl.append(p_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abs_id = OxmlElement("w:abstractNumId")
    abs_id.set(qn("w:val"), str(abstract_id))
    num.append(abs_id)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)


def add_bullet(doc, text: str, *, bold_prefix: str | None = None):
    p = doc.add_paragraph()
    apply_numbering(p, doc._bullet_num_id)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)
    set_paragraph_keep(p)
    return p


def add_numbered(doc, text: str):
    p = doc.add_paragraph()
    apply_numbering(p, doc._decimal_num_id)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    set_run_font(r)
    set_paragraph_keep(p)
    return p


def add_body(doc, text: str, *, bold_prefix: str | None = None, color=INK):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, bold=True, color=color)
        r = p.add_run(text[len(bold_prefix):])
        set_run_font(r, color=color)
    else:
        r = p.add_run(text)
        set_run_font(r, color=color)
    set_paragraph_keep(p)
    return p


def add_code_line(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.18)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    set_run_font(r, name=CODE_FONT, size=8.5, color=NAVY)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), LIGHT_GRAY)
    p._p.get_or_add_pPr().append(shd)
    set_paragraph_keep(p)
    return p


def add_callout(doc, label: str, text: str, *, fill=PALE_BLUE, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, color=accent, size="8")
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(f"{label}  ")
    set_run_font(r, bold=True, color=accent)
    r = p.add_run(text)
    set_run_font(r, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_table(doc, headers, rows, widths_dxa, *, font_size=9.2, header_fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths_dxa)
    set_repeat_table_header(table.rows[0])
    set_table_row_cant_split(table.rows[0])
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, header_fill)
        set_cell_border(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.05
        r = p.add_run(str(header))
        header_text_color = WHITE if header_fill in (NAVY, DARK_BLUE) else NAVY
        set_run_font(r, size=font_size, bold=True, color=header_text_color)
    for row_data in rows:
        row = table.add_row()
        set_table_row_cant_split(row)
        for i, value in enumerate(row_data):
            cell = row.cells[i]
            set_cell_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.08
            r = p.add_run(str(value))
            set_run_font(r, size=font_size, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_heading(doc, text: str, level: int = 1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    set_paragraph_keep(p, keep_next=True)
    return p


def add_page_number(paragraph):
    run = paragraph.add_run("Page ")
    set_run_font(run, size=8.5, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def configure_styles(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    normal.font.name = BASE_FONT
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal._element.rPr.rFonts.set(qn("w:ascii"), BASE_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), BASE_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = doc.styles["Title"]
    title.font.name = BASE_FONT
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = rgb(NAVY)
    title._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(8)

    specs = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in specs.items():
        style = doc.styles[name]
        style.font.name = BASE_FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), BASE_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), BASE_FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header
    p = header.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    r = p.add_run("CANADA LOGISTICS MARKETPLACE")
    set_run_font(r, size=8.5, bold=True, color=MUTED)
    r = p.add_run("\tARCHITECTURE BASELINE v1.0")
    set_run_font(r, size=8.5, color=MUTED)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    add_page_number(p)

    doc._bullet_num_id = add_numbering_definition(doc, kind="bullet")
    doc._decimal_num_id = add_numbering_definition(doc, kind="decimal")


def font(size, bold=False):
    path = "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf"
    return ImageFont.truetype(path, size=size)


def draw_centered(draw, box, text, fnt, fill):
    lines = text.split("\n")
    heights = []
    widths = []
    for line in lines:
        b = draw.textbbox((0, 0), line, font=fnt)
        widths.append(b[2] - b[0])
        heights.append(b[3] - b[1])
    total_h = sum(heights) + (len(lines) - 1) * 8
    y = box[1] + (box[3] - box[1] - total_h) / 2
    for line, w, h in zip(lines, widths, heights):
        x = box[0] + (box[2] - box[0] - w) / 2
        draw.text((x, y), line, font=fnt, fill=fill)
        y += h + 8


def draw_box(draw, box, text, *, fill, outline=BLUE, text_color=NAVY, title=False):
    fill_value = fill if fill.startswith("#") else "#" + fill
    outline_value = outline if outline.startswith("#") else "#" + outline
    text_value = text_color if text_color.startswith("#") else "#" + text_color
    draw.rounded_rectangle(box, radius=18, fill=fill_value, outline=outline_value, width=3)
    draw_centered(draw, box, text, font(29 if title else 25, bold=title), text_value)


def draw_arrow(draw, start, end, *, color=MUTED, width=5):
    color_tuple = "#" + color
    draw.line([start, end], fill=color_tuple, width=width)
    ex, ey = end
    sx, sy = start
    if abs(ex - sx) >= abs(ey - sy):
        sign = 1 if ex > sx else -1
        pts = [(ex, ey), (ex - sign * 18, ey - 11), (ex - sign * 18, ey + 11)]
    else:
        sign = 1 if ey > sy else -1
        pts = [(ex, ey), (ex - 11, ey - sign * 18), (ex + 11, ey - sign * 18)]
    draw.polygon(pts, fill=color_tuple)


def build_domain_flow(path: Path):
    img = Image.new("RGB", (1800, 760), "#" + WHITE)
    draw = ImageDraw.Draw(img)
    draw.text((70, 34), "Marketplace domain flow", font=font(38, True), fill="#" + NAVY)
    top = [
        ("Shipment\nRequest", PALE_BLUE),
        ("Planning", PALE_BLUE),
        ("RFQ", LIGHT_BLUE),
        ("Supplier\nQuote", LIGHT_BLUE),
        ("Quote\nCenter", "EAF4EF"),
        ("Order", "EAF4EF"),
    ]
    x0, y0, w, h, gap = 60, 130, 235, 130, 48
    boxes = []
    for i, (label, fill) in enumerate(top):
        box = (x0 + i * (w + gap), y0, x0 + i * (w + gap) + w, y0 + h)
        boxes.append(box)
        draw_box(draw, box, label, fill="#" + fill, title=True)
        if i:
            draw_arrow(draw, (boxes[i - 1][2], y0 + h / 2), (box[0], y0 + h / 2))
    bottom = [
        ("Matching", LIGHT_BLUE),
        ("Fulfillment", LIGHT_BLUE),
        ("Tracking", PALE_BLUE),
        ("Supplier\nCenter", PALE_BLUE),
    ]
    bx0, by0 = 310, 390
    bboxes = []
    for i, (label, fill) in enumerate(bottom):
        box = (bx0 + i * (w + 95), by0, bx0 + i * (w + 95) + w, by0 + h)
        bboxes.append(box)
        draw_box(draw, box, label, fill="#" + fill, title=True)
        if i:
            draw_arrow(draw, (bboxes[i - 1][2], by0 + h / 2), (box[0], by0 + h / 2))
    draw_arrow(draw, ((boxes[-1][0] + boxes[-1][2]) / 2, boxes[-1][3]), (bboxes[0][0] + w / 2, bboxes[0][1]))
    pi = (460, 585, 1340, 730)
    draw_box(draw, pi, "PRICE INTELLIGENCE\nCollector → Forecast → Recommendation → Dynamic Pricing", fill="#FFF8E8", outline=AMBER, text_color=AMBER, title=False)
    draw_arrow(draw, (boxes[3][0] + w / 2, boxes[3][3]), (700, pi[1]))
    draw_arrow(draw, (bboxes[0][0] + w / 2, bboxes[0][3]), (900, pi[1]))
    draw_arrow(draw, (1200, pi[1]), (boxes[4][0] + w / 2, boxes[4][3]))
    img.save(path)


def build_runtime_arch(path: Path):
    img = Image.new("RGB", (1800, 980), "#" + WHITE)
    draw = ImageDraw.Draw(img)
    draw.text((70, 34), "MVP runtime architecture", font=font(38, True), fill="#" + NAVY)
    lane_titles = ["CLIENTS", "APPLICATION PLANE", "STATE PLANE", "EXTERNAL SERVICES"]
    lane_x = [60, 430, 890, 1340]
    lane_w = [300, 390, 380, 390]
    for title, x, lw in zip(lane_titles, lane_x, lane_w):
        draw.rounded_rectangle((x, 105, x + lw, 900), radius=22, fill="#F8FAFC", outline="#D8E0EA", width=3)
        draw.text((x + 24, 130), title, font=font(24, True), fill="#" + MUTED)
    clients = [(110, 230, 310, 340, "Browser"), (110, 425, 310, 535, "Supplier API"), (110, 620, 310, 730, "Webhooks")]
    for x1, y1, x2, y2, label in clients:
        draw_box(draw, (x1, y1, x2, y2), label, fill="#" + PALE_BLUE, title=True)
    apps = [(500, 210, 750, 320, "Next.js\nWeb"), (500, 405, 750, 535, "NestJS\nAPI"), (500, 630, 750, 760, "NestJS\nWorker")]
    for x1, y1, x2, y2, label in apps:
        draw_box(draw, (x1, y1, x2, y2), label, fill="#" + LIGHT_BLUE, title=True)
    state = [(955, 190, 1205, 295, "PostgreSQL"), (955, 345, 1205, 450, "RabbitMQ"), (955, 500, 1205, 605, "Redis"), (955, 655, 1205, 760, "S3 Storage")]
    for x1, y1, x2, y2, label in state:
        draw_box(draw, (x1, y1, x2, y2), label, fill="#EAF4EF", outline=GREEN, text_color=GREEN, title=True)
    external = [(1410, 190, 1660, 295, "Google Maps"), (1410, 345, 1660, 450, "Resend"), (1410, 500, 1660, 605, "AI Gateway"), (1410, 655, 1660, 760, "Supplier\nEndpoints")]
    for x1, y1, x2, y2, label in external:
        draw_box(draw, (x1, y1, x2, y2), label, fill="#FFF8E8", outline=AMBER, text_color=AMBER, title=True)
    draw_arrow(draw, (310, 285), (500, 265))
    draw_arrow(draw, (310, 480), (500, 470))
    draw_arrow(draw, (310, 675), (500, 470))
    draw_arrow(draw, (625, 320), (625, 405))
    draw_arrow(draw, (750, 470), (955, 245))
    draw_arrow(draw, (750, 470), (955, 397))
    draw_arrow(draw, (750, 695), (955, 397))
    draw_arrow(draw, (750, 695), (955, 552))
    draw_arrow(draw, (750, 695), (955, 707))
    draw_arrow(draw, (750, 470), (1410, 242))
    draw_arrow(draw, (750, 695), (1410, 397))
    draw_arrow(draw, (750, 695), (1410, 552))
    draw_arrow(draw, (750, 695), (1410, 707))
    draw.text((500, 835), "PostgreSQL is the system of record  ·  Redis is disposable  ·  External work is asynchronous", font=font(24, True), fill="#" + NAVY)
    img.save(path)


def add_picture_with_alt(doc, path: Path, alt: str, width=6.5):
    shape = doc.add_picture(str(path), width=Inches(width))
    shape._inline.docPr.set("descr", alt)
    p = doc.paragraphs[-1]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    return shape


def build_document():
    domain_image = QA_DIR / "domain_flow.png"
    runtime_image = QA_DIR / "runtime_architecture.png"
    build_domain_flow(domain_image)
    build_runtime_arch(runtime_image)

    doc = Document()
    configure_styles(doc)

    # Cover / memo masthead
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(26)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("ARCHITECTURE BASELINE")
    set_run_font(r, size=10, bold=True, color=BLUE)

    p = doc.add_paragraph("Canada Logistics Marketplace", style="Title")
    p.paragraph_format.space_after = Pt(4)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(22)
    r = p.add_run("加拿大物流服务交易平台 · MVP 目标架构与实施基线")
    set_run_font(r, size=14, color=MUTED)

    add_table(
        doc,
        ["项目元数据", "值"],
        [
            ("版本", "v1.0"),
            ("日期", "2026-07-17"),
            ("架构形态", "Next.js + NestJS 模块化单体 + 独立 Worker"),
            ("数据/消息", "PostgreSQL + Redis + RabbitMQ + S3 Compatible"),
            ("范围", "目录、ER、模块关系、API、DDL、技术架构、MVP 计划"),
            ("明确不包含", "业务代码、Prisma migration、现有系统覆盖或删除"),
        ],
        [2050, 7310],
        font_size=10,
    )
    add_callout(doc, "架构结论", "MVP 先做模块化单体和可靠异步，不提前拆微服务；交易的是可组合物流服务，AI 只产生候选与建议，确定性校验和审批负责放行。")
    add_table(
        doc,
        ["72", "4", "15", "7"],
        [("数据库表", "只读视图", "限界上下文", "正式架构交付项")],
        [2340, 2340, 2340, 2340],
        font_size=10,
        header_fill=NAVY,
    )

    doc.add_page_break()
    add_heading(doc, "0. 阅读导航", 1)
    add_body(doc, "本 Word 文档用于架构评审；仓库中的 Markdown 与 SQL 是可追踪的详细基线。")
    add_table(
        doc,
        ["章节", "回答的问题", "详细产物"],
        [
            ("1 目标架构与目录", "代码如何按 DDD 落地？", "01-directory-and-ddd.md"),
            ("2 数据模型与 ER", "核心实体、版本、约束是什么？", "02-data-model-and-er.md + DDL"),
            ("3 模块关系", "同步/异步边界和事件是什么？", "03-module-context-map.md"),
            ("4 API", "四角色与外部系统如何接入？", "04-api-design.md"),
            ("5 技术架构", "如何运行、扩缩、观测和保护？", "05-technical-architecture.md"),
            ("6 MVP 计划", "按什么顺序开发和验收？", "06-mvp-delivery-plan.md"),
            ("7 ADR/评审门", "哪些选择已定、哪些待业务确认？", "07-architecture-decisions.md"),
        ],
        [1150, 3900, 4310],
    )

    add_heading(doc, "执行摘要", 1)
    for text in [
        "服务能力而非线路产品：服务目录、Offering、Coverage、Plan Leg 和 Fulfillment Task 都是数据，可在后台扩展。",
        "API First：Web、供应商门户、Email/API/Excel/Webhook 全部进入同一应用层合约。",
        "强事务 + 可靠异步：同步命令只提交 PostgreSQL 事务和 Outbox；邮件、AI、Webhook、Forecast 由 Worker 执行。",
        "组织级 RBAC：用户通过 Membership 在不同组织获得角色；供应商只可访问本组织 RFQ、报价和任务。",
        "不可变版本：Plan、Supplier Quote、Customer Quote 修改产生新 revision，客户接受与履约引用当时快照。",
        "Price Intelligence 独立：Collector、Forecast、Recommendation、Dynamic Pricing 不直接拥有或改写 Quote。",
        "Legacy 安全迁移：现有 FastAPI/Vite 与确定性 Zone Engine 先保留，以黄金样例、dual-run 和 Feature Flag 迁移。",
    ]:
        add_bullet(doc, text)

    add_heading(doc, "1. 目标目录与 DDD 模块", 1)
    add_code_line(doc, "apps/web        Next.js App Router；四角色门户")
    add_code_line(doc, "apps/api        NestJS HTTP API；模块化单体")
    add_code_line(doc, "apps/worker     RabbitMQ consumer、scheduler、outbox relay")
    add_code_line(doc, "packages/       contracts / database / config / security / observability / ui")
    add_code_line(doc, "docs/           architecture / api / domains / runbooks")
    add_code_line(doc, "infra/          docker / kubernetes / terraform / monitoring")
    add_code_line(doc, "apps/legacy-quote-v1  迁移期保留的现有 FastAPI/Vite 能力")
    add_body(doc, "每个 NestJS 领域模块使用统一结构：Domain（Entity/Value Object/Event/Policy/Repository Port）→ Application（Command/Query/Service/DTO/Port）→ Infrastructure（Prisma/Messaging/Integration Adapter）→ Interface（Controller/Validation/Presenter/Consumer）。")

    add_heading(doc, "模块所有权", 2)
    add_table(
        doc,
        ["模块", "核心职责", "拥有的数据"],
        [
            ("Identity & Access", "Organization、Membership、JWT、RBAC、API Client", "users / organizations / roles / sessions"),
            ("Service Catalog", "服务、关系、Offering、Coverage、Facility", "service_* / supplier_service_* / facilities"),
            ("Shipment Request", "地址、Cargo、客户服务需求", "locations / shipment_requests / cargo / services"),
            ("Planning", "AI/规则 Plan、DAG 校验、revision、审批", "logistics_plans / plan_legs / dependencies"),
            ("RFQ", "RFQ 轮次、选商、邀请、渠道发送", "rfqs / items / invitations / dispatch attempts"),
            ("Supplier Quote", "回复接入、AI 解析、标准化、复核", "supplier_quotes / items / charges / parse runs"),
            ("Quote Center", "最低/最快/推荐/DDP/Self Import", "customer_quotes / options / acceptances"),
            ("Matching + Order", "拼货池、锁价、订单、Shipment 编组", "pools / locks / orders / shipments"),
            ("Fulfillment + Tracking", "任务拆分、依赖、状态、统一 Timeline", "tasks / history / milestones / events"),
            ("Supplier Center", "档案、资质、投诉、KPI", "supplier_profiles / complaints / kpi snapshots"),
            ("Price Intelligence", "收集、预测、推荐、动态定价", "observations / forecasts / policies / decisions"),
            ("Integration + AI Governance", "文件、连接、消息、模型证据、复核", "assets / outbox / inbox / ai runs / reviews"),
        ],
        [1850, 3960, 3550],
        font_size=8.7,
    )
    add_callout(doc, "依赖规则", "Domain 不依赖 NestJS/Prisma；跨模块不得直接访问对方 Repository，只能调用公开 Application Facade 或消费版本化领域事件。")

    add_heading(doc, "2. 模块关系与主流程", 1)
    add_picture_with_alt(doc, domain_image, "Marketplace domain flow showing request, planning, RFQ, quotes, order, matching, fulfillment, tracking, supplier center and price intelligence.")
    p = doc.add_paragraph("图 1 · Marketplace 领域主流程与 Price Intelligence 反馈回路")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    for run in p.runs:
        set_run_font(run, size=9, italic=True, color=MUTED)
    add_body(doc, "主交易链路：客户提交 Shipment Request → Hybrid Planner 生成并校验 Plan → RFQ 选商与多渠道分发 → 供应商回复/AI 解析 → 多方案 Customer Quote → 接受后生成 Order → 拼货/锁价 → Shipment/Task → Tracking。")

    add_heading(doc, "关键领域事件", 2)
    add_table(
        doc,
        ["事件", "Producer", "主要 Consumer"],
        [
            ("ShipmentRequestSubmitted.v1", "Shipment Request", "Planning / Audit"),
            ("LogisticsPlanApproved.v1", "Planning", "RFQ / Shipment Request"),
            ("RfqPublished.v1", "RFQ", "Integration / Supplier Center"),
            ("RfqReplyReceived.v1", "Integration", "Supplier Quote"),
            ("SupplierQuoteValidated.v1", "Supplier Quote", "Quote Center / Price Intelligence"),
            ("CustomerQuoteAccepted.v1", "Quote Center", "Order / Price Intelligence"),
            ("PoolCapacityChanged.v1", "Matching", "Price Intelligence / Dashboard"),
            ("ShipmentAllocated.v1", "Order/Matching", "Fulfillment / Tracking"),
            ("FulfillmentTaskStatusChanged.v1", "Fulfillment", "Tracking / Supplier KPI / Notification"),
            ("PriceForecastGenerated.v1", "Price Intelligence", "Quote Center / Dashboard"),
        ],
        [3800, 2300, 3260],
        font_size=8.8,
    )

    add_heading(doc, "3. 数据模型与 ER 基线", 1)
    add_callout(doc, "基线规模", "PostgreSQL marketplace schema：72 张表、4 个可重建只读视图；所有核心交易均有组织归属、版本、来源和审计路径。")
    add_table(
        doc,
        ["数据域", "核心实体", "核心关系/目的"],
        [
            ("IAM", "User, Organization, Membership, Role, Permission", "一个用户多组织；权限按当前组织计算"),
            ("Catalog/Supplier", "Service, Relationship, Offering, Coverage, Facility", "服务动态扩展；供应商能力可筛选"),
            ("Request/Plan", "Shipment Request, Cargo, Plan, Plan Leg", "客户意图 → 可采购/可履约 DAG"),
            ("RFQ/Supplier Quote", "RFQ, Invitation, Dispatch, Quote, Charge", "多轮/多供应商/多渠道；原始回复可追溯"),
            ("Customer Quote", "Quote, Option, Option Leg, Acceptance", "成本组合、比较、发布和不可变接受"),
            ("Matching/Order", "Pool, Membership, Lock, Order, Shipment", "并发容量、价格锁、多单拼货/未来拆单"),
            ("Fulfillment/Tracking", "Task, Dependency, History, Milestone, Event", "供应商任务隔离与统一 Timeline"),
            ("Price Intelligence", "Observation, Model, Forecast, Recommendation, Policy, Decision", "有证据的价格大脑，不直接改写报价"),
            ("Platform", "File, Connection, AI Run, Review, Outbox, Inbox, Audit", "可靠异步、外部集成、模型治理和审计"),
        ],
        [1850, 3850, 3660],
        font_size=8.8,
    )

    add_heading(doc, "数据约定", 2)
    conventions = [
        ("ID", "UUID；不把数据库自增 ID 暴露为领域标识"),
        ("时间", "timestamptz；API 使用 ISO 8601 UTC"),
        ("金额", "numeric(18,4) + ISO currency；API 以字符串传输"),
        ("价格分析", "原币金额 + CAD 归一化金额 + FX snapshot"),
        ("重量/体积", "数据库统一 kg/cbm；API 边界换算"),
        ("状态", "text + CHECK；重要 Aggregate 使用 version CAS"),
        ("JSONB", "外部 payload、快照、扩展条件；可检索核心字段必须结构化"),
        ("删除", "Master Data 可逻辑删；交易/事件/审计追加式保留"),
    ]
    add_table(doc, ["主题", "约定"], conventions, [1900, 7460], font_size=9.2)

    add_heading(doc, "数据库必须保证的关键不变量", 2)
    for text in [
        "RFQ 和 Customer Quote 所引用的 Plan 必须属于同一 Shipment Request。",
        "Supplier Quote 的 Supplier 必须与 RFQ Invitation 的 Supplier 相同。",
        "被接受的 Quote Option 必须属于该 Customer Quote；Order 必须引用同一 Acceptance/Option。",
        "同一 RFQ 对同一 Supplier 只允许一条 Invitation；报价/计划/客户报价 revision 唯一。",
        "一个 Customer Quote 最多一个 Recommended Option。",
        "Tracking Event、Price Observation、Task Status History、Supplier KPI、Audit Log 禁止 UPDATE/DELETE。",
        "HTTP command、Outbox/Inbox、Webhook、通知和外部 Tracking Event 均有去重/幂等键。",
    ]:
        add_bullet(doc, text)
    add_body(doc, "Plan DAG 无环、服务完整覆盖、Quote Option 成本腿完整、Pool 并发不超容量、Task 合法状态转换等复杂不变量由 Domain + 真实 PostgreSQL 集成测试保证。")

    add_heading(doc, "4. REST API 设计", 1)
    add_table(
        doc,
        ["协议主题", "约定"],
        [
            ("Base", "/api/v1；OpenAPI 是前端/外部集成唯一合约"),
            ("Auth", "短时 JWT；Refresh rotation；API Client secret 只显示一次"),
            ("Org Context", "X-Organization-Id + Membership/Permission + Resource Ownership"),
            ("Idempotency", "创建/发布/接受/分配/状态命令必须 Idempotency-Key"),
            ("Concurrency", "If-Match: <version>；冲突返回 409"),
            ("Async", "长任务返回 202 + operationId；GET /operations/{id}"),
            ("Errors", "Problem Details + 稳定业务 code + field errors + requestId"),
            ("Pagination", "Cursor；默认 25、最大 100"),
            ("Money", "十进制字符串；永不使用 JSON 浮点计算金额"),
        ],
        [1900, 7460],
        font_size=9.1,
    )

    add_heading(doc, "API Surface", 2)
    add_table(
        doc,
        ["Surface", "代表端点", "主要角色"],
        [
            ("Auth/RBAC", "POST /auth/login · GET /me · /organizations · /roles", "All / Admin / Org Owner"),
            ("Catalog", "GET/POST /services · /supplier-offerings · /facilities", "Customer / Supplier / Admin"),
            ("Request/Plan", "POST /shipment-requests · /submit · /logistics-plans · /approve", "Customer / Sales / Admin"),
            ("RFQ", "POST /rfqs · /supplier-candidates · /publish · /resend", "Sales / Admin"),
            ("Supplier Quote", "/reply-files · /quote-parse-runs · /resolve · /supplier/.../quotes", "Supplier / Sales / Admin"),
            ("Quote Center", "/customer-quotes · /comparison · /publish · /acceptances", "Customer / Sales / Admin"),
            ("Matching/Order", "/orders · /consolidation-pools · /memberships · /price-locks", "Sales / Admin / System"),
            ("Fulfillment", "/shipments · /supplier/fulfillment-tasks · /complete · /exceptions", "Supplier / Sales / Admin"),
            ("Tracking", "/shipments/{id}/timeline · /tracking-events · /webhooks/inbound", "All scoped / External signed"),
            ("Price Intelligence", "/forecasts · /recommendations · /policies · /pricing-decisions", "Sales / Analyst / Admin"),
            ("Platform", "/files · /integrations · /review-tasks · /audit-logs · /dashboards", "Scoped / Admin / Auditor"),
        ],
        [1800, 5000, 2560],
        font_size=8.7,
    )
    add_callout(doc, "权限边界", "Controller 先验证 JWT/Permission，再由 Repository/Application Policy 强制 customer/supplier ownership。供应商 Task 查询始终带 current supplier organization filter。", fill="FFF8E8", accent=AMBER)

    add_heading(doc, "5. 技术架构", 1)
    add_picture_with_alt(doc, runtime_image, "MVP runtime architecture with clients, Next.js, NestJS API and worker, PostgreSQL, RabbitMQ, Redis, S3, Maps, Resend, AI gateway and supplier endpoints.")
    p = doc.add_paragraph("图 2 · MVP Runtime：同步事务与异步 Worker 分离")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    for run in p.runs:
        set_run_font(run, size=9, italic=True, color=MUTED)

    add_heading(doc, "运行与可靠性", 2)
    for text in [
        "API 同步链路不等待 Email、Webhook、AI 或 Forecast；事务内写业务状态、Audit 和 Outbox。",
        "Outbox Relay 用 SKIP LOCKED 领取，RabbitMQ 使用持久消息/Publisher Confirm；Consumer 用 Inbox 去重。",
        "Redis 仅做缓存、限流和短锁；价格、容量、Order/Task 状态最终以 PostgreSQL 为准。",
        "S3 保存原始回复、Excel/PDF、POD 和模型 artifact；数据库保存对象引用、SHA-256、扫描状态和归属。",
        "Queue 独立配置并发、timeout、重试、退避和 DLQ；业务错误进入 Review，不做无效重试。",
    ]:
        add_bullet(doc, text)

    add_heading(doc, "初始 SLO", 2)
    add_table(
        doc,
        ["能力", "目标"],
        [
            ("核心读 API", "99.9% 月可用；P95 < 500 ms（不含 AI/文件）"),
            ("核心写 API", "P95 < 800 ms；事务不等待外部服务"),
            ("Outbox", "P95 事件年龄 < 5 s"),
            ("RFQ 分发", "P95 开始发送 < 60 s"),
            ("AI Quote Parse", "P95 < 3 min；失败可重试/人工"),
            ("Tracking", "P95 事件可见 < 30 s"),
            ("安全/正确性", "重复副作用、跨组织越权、无来源金额发布均为 0"),
            ("DLQ", "未告警时间 < 5 min"),
        ],
        [2800, 6560],
    )

    add_heading(doc, "6. AI 与 Price Intelligence", 1)
    add_heading(doc, "AI 安全边界", 2)
    add_table(
        doc,
        ["能力", "AI 可以做", "AI 不可以做 / 放行者"],
        [
            ("Logistics Plan", "返回结构化候选 Plan", "不能批准；Schema/DAG/服务/地点/能力 validator 放行"),
            ("Quote Parse", "提取 price/currency/THC/DOC/transit/free time/notes", "不能绕过来源/算术/RFQ scope；Review/validator 放行"),
            ("Forecast", "产生有版本、有区间的预测", "不能隐藏训练截止/样本/误差；Registry + Shadow 放行"),
            ("Recommendation", "建议销售价/采购价区间", "不能直接更新 Quote；Policy Guard + Approver 放行"),
            ("Explanation", "解释已锁定结果与风险", "不能生成未在证据中的金额"),
        ],
        [1700, 3730, 3930],
        font_size=8.7,
    )

    add_heading(doc, "Price Intelligence 四个子模块", 2)
    add_table(
        doc,
        ["子模块", "输入", "持久化/输出", "关键护栏"],
        [
            ("Collector", "Supplier Quote /成交/价卡/历史", "Price Observation + lineage", "原币/CAD/FX/权限/质量；追加修正"),
            ("Forecast", "Time-safe segment series", "7/14/30 天 P10/P50/P90", "训练截止、样本数、回测、Candidate→Shadow→Active"),
            ("Recommendation", "成本/Forecast/容量/热度", "Sales/Procurement 区间", "低置信/过期不可用；完整 evidence"),
            ("Dynamic Pricing", "Recommendation + versioned policy", "Audited Pricing Decision", "毛利 floor、折扣/加价 cap、人工阈值、new quote revision"),
        ],
        [1600, 2450, 2600, 2710],
        font_size=8.6,
    )
    add_callout(doc, "核心竞争力", "Price Intelligence 消费所有合法交易事实，但不反向拥有 RFQ/Quote/Order。它通过 Advisory Facade 和版本化事件提供建议，保持可解释、可回测、可审批和可回退。")

    add_heading(doc, "7. MVP 开发计划", 1)
    add_body(doc, "完整 PRD 不是小型 CRUD。以 1 PO、1 Tech Lead、2 Backend、2 Frontend、1 Data/ML、1 QA（DevOps/Security 共享）为基线：Pilot 核心闭环约 14-16 周；完整 MVP 约 22-26 周。")
    add_table(
        doc,
        ["Phase", "范围", "主要 Exit Gate"],
        [
            ("0 架构", "目录/ER/关系/API/DDL/技术/计划", "Product/Tech/Data/Security 评审"),
            ("1 数据库", "Prisma schema、baseline migration、seed、Outbox/Inbox", "空库/升级迁移、约束负向测试、恢复演练"),
            ("2 IAM", "用户、组织、RBAC、JWT、API Client", "四角色 E2E；跨组织越权为 0"),
            ("3 Catalog/Plan", "服务市场、Request、Hybrid Planner", "示例流程可生成；无效 Plan 不批准"),
            ("4 RFQ", "选商、Email/API/Excel/Webhook", "批量幂等；失败可重试/DLQ"),
            ("5 Supplier/Parse", "档案、回复接入、AI 解析、Review", "金额/币种错误 0 次自动放行"),
            ("6 Quote/Price Core", "多方案、Collector、Recommendation", "成本完整；无来源金额不可发布；Pilot Gate"),
            ("7 Matching", "Pool、容量、锁价、通知", "并发无超卖；锁/Reservation 自动释放"),
            ("8 Fulfillment", "Order/Shipment/Task/dependency", "合法状态机；供应商任务隔离"),
            ("9 Tracking", "Milestone/Event/Timeline/POD", "重复/乱序安全；客户字段脱敏"),
            ("10 Intelligence", "Dashboard/KPI/Forecast/Dynamic Pricing", "Shadow/误差/漂移/审批/Guardrail 全通过"),
        ],
        [1100, 3910, 4350],
        font_size=8.5,
    )

    add_heading(doc, "Phase 1 数据库开发的明确边界", 2)
    for text in [
        "新增 packages/database；将 baseline DDL 翻译为 Prisma schema 和 0001 migration。",
        "建立平台组织、系统角色/权限、初始服务/关系的可重复 Seed。",
        "使用真实 PostgreSQL 执行空库迁移、上一版本升级、constraint/lock/trigger/view 测试。",
        "建立 schema drift 和 destructive migration gate；不得使用 db push 代替 migration。",
        "保留现有 FastAPI/Vite；数据库阶段不删除、不覆盖、不直接导入生产价卡。",
    ]:
        add_numbered(doc, text)

    add_heading(doc, "8. 关键 ADR、风险与评审门", 1)
    add_table(
        doc,
        ["ADR", "决策"],
        [
            ("001", "模块化单体 + Worker；有组织/性能/合规证据后再拆服务"),
            ("002", "PostgreSQL 是唯一事实源；Redis 可丢失"),
            ("003", "一个 marketplace schema；逻辑表所有权"),
            ("004", "组织级 Membership/Role/Permission"),
            ("005", "Plan/Supplier Quote/Customer Quote 不可变 revision"),
            ("006", "Plan 使用 DAG；界面可线性展示"),
            ("007", "AI 候选/建议；确定性 Guardrail/审批放行"),
            ("008", "原币事实 + CAD/FX 分析快照"),
            ("009", "Outbox/Inbox + RabbitMQ 至少一次投递"),
            ("010", "REST/OpenAPI First；GraphQL 暂不实现"),
            ("011", "Order 与 Shipment 分离，支持拼货/拆单"),
            ("012-014", "Price Intelligence 独立且保护价格隐私；Legacy Strangler 迁移"),
        ],
        [1050, 8310],
        font_size=9.0,
    )

    add_heading(doc, "主要风险", 2)
    add_table(
        doc,
        ["风险", "主要缓解"],
        [
            ("服务组合不完整", "Catalog relationship + deterministic validator + SME review"),
            ("回复格式异构", "Channel extractor + evidence + gold set + confidence gate"),
            ("币种/费用口径", "原币 + FX snapshot + charge taxonomy + arithmetic gate"),
            ("跨组织越权", "Org RBAC + ownership filter + BOLA tests + audit"),
            ("消息重复", "Outbox/Inbox + idempotency + unique constraints"),
            ("拼货超卖", "PostgreSQL row lock + version CAS + concurrency tests"),
            ("Forecast 小样本/漂移", "样本门槛 + 区间 + Shadow + drift + fallback"),
            ("Dynamic Pricing 亏损", "margin floor/cap + approval + revision + audit"),
            ("Legacy 回归", "golden fixtures + dual-run + feature flag"),
        ],
        [3000, 6360],
        font_size=9.0,
    )

    add_heading(doc, "架构评审门", 2)
    for text in [
        "同意模块化单体 + 独立 Worker。",
        "同意组织级 RBAC 和供应商任务隔离。",
        "同意 marketplace schema、72 表/4 视图 baseline。",
        "同意不可变 revision、Order/Shipment 分离、Outbox/Inbox。",
        "同意 AI/Price Intelligence 的证据、Shadow、Guardrail 和审批边界。",
        "同意 Legacy 在数据库阶段继续保留。",
        "确认 Phase 1 只做数据库、迁移、Seed 和基础设施测试，不写业务功能。",
    ]:
        add_bullet(doc, text)

    add_heading(doc, "进入实施前仍需业务确认", 2)
    add_table(
        doc,
        ["事项", "当前默认", "Owner / 截止"],
        [
            ("FX 来源与锁定时点", "每次 Observation/Quote 保存来源和快照", "Finance/Product · Phase 5"),
            ("GST/CARM/DDP 责任", "作为服务，不自动做法律/税务判断", "Customs/Finance · Phase 3"),
            ("Supplier KPI/推荐权重", "指标分开存、权重版本化", "Procurement/Sales · Phase 4-6"),
            ("Forecast 样本/误差门槛", "低样本不出细粒度预测", "Data/Product · Phase 6"),
            ("聚合价格匿名门槛", "未定前不对外显示", "Security/Product · Phase 6"),
            ("Dynamic Pricing 自动阈值", "MVP 外部发布默认人工批准", "Sales/Finance · Phase 10"),
            ("保留期限/AI Provider", "架构建议 + Provider-neutral port", "Legal/Security · 上线前"),
        ],
        [2820, 3870, 2670],
        font_size=8.8,
    )

    add_heading(doc, "附录：仓库产物", 1)
    add_table(
        doc,
        ["产物", "路径"],
        [
            ("架构索引", "docs/architecture/README.md"),
            ("目录与 DDD", "docs/architecture/01-directory-and-ddd.md"),
            ("数据模型与 ER", "docs/architecture/02-data-model-and-er.md"),
            ("模块 Context Map", "docs/architecture/03-module-context-map.md"),
            ("REST API", "docs/architecture/04-api-design.md"),
            ("技术架构", "docs/architecture/05-technical-architecture.md"),
            ("MVP 计划", "docs/architecture/06-mvp-delivery-plan.md"),
            ("ADR/评审门", "docs/architecture/07-architecture-decisions.md"),
            ("PostgreSQL DDL", "database/marketplace_v1.sql"),
        ],
        [2600, 6760],
        font_size=9.3,
    )
    add_callout(doc, "验证状态", "Markdown 相对链接/代码围栏、DDL 表/视图数量、外键目标、建表顺序、对象重名、触发表和括号平衡已通过静态检查。真实 PostgreSQL 空库执行列为 Phase 1 第一项，因为当前工作环境没有可用 PostgreSQL 服务。", fill=LIGHT_GRAY, accent=MUTED)

    core_props = doc.core_properties
    core_props.title = "Canada Logistics Marketplace Architecture Baseline v1.0"
    core_props.subject = "MVP architecture, DDD, ER, API, DDL, technical design and delivery plan"
    core_props.author = "Architecture Team"
    core_props.keywords = "logistics marketplace, DDD, NestJS, Next.js, PostgreSQL, RFQ, price intelligence"

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_document()
